from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# --- Margens ---
sec = doc.sections[0]
sec.top_margin    = Cm(2.5)
sec.bottom_margin = Cm(2.5)
sec.left_margin   = Cm(3)
sec.right_margin  = Cm(3)

# --- Cores ---
GREEN_DARK = RGBColor(0x1A, 0x5C, 0x4A)
GREEN_MID  = RGBColor(0x2E, 0x8B, 0x6A)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_TEXT  = RGBColor(0x44, 0x44, 0x44)
BLACK      = RGBColor(0x11, 0x11, 0x11)

def style_run(run, size=10.5, color=GRAY_TEXT, bold=False, italic=False):
    run.font.size      = Pt(size)
    run.font.color.rgb = color
    run.font.bold      = bold
    run.font.italic    = italic
    run.font.name      = 'Calibri'

def add_h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text.upper())
    style_run(r, size=13, color=GREEN_DARK, bold=True)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '4');    bot.set(qn('w:color'), '1A5C4A')
    pBdr.append(bot); pPr.append(pBdr)

def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    style_run(r, size=11.5, color=GREEN_MID, bold=True)

def add_h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    style_run(r, size=10.5, color=BLACK, bold=True)

def add_body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    style_run(r)

def add_kv(key, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    rk = p.add_run(key + ': ')
    style_run(rk, bold=True, color=BLACK)
    rv = p.add_run(value)
    style_run(rv)

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.5)
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        style_run(rb, bold=True, color=BLACK)
    r = p.add_run(text)
    style_run(r)

def add_table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        c = hrow.cells[i]
        c.text = h
        r = c.paragraphs[0].runs[0]
        style_run(r, size=10, color=WHITE, bold=True)
        tc = c._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1A5C4A'); tcPr.append(shd)
    for ri, row in enumerate(rows):
        drow = t.rows[ri+1]
        fill = 'F5F5F5' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            c = drow.cells[ci]
            c.text = val
            r = c.paragraphs[0].runs[0]
            style_run(r, size=10, color=GRAY_TEXT)
            tc = c._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill); tcPr.append(shd)
    doc.add_paragraph()

def add_sep():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '4')
    bot.set(qn('w:space'), '1');    bot.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bot); pPr.append(pBdr)

# ============================================================
# CAPA
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
r = p.add_run('BASE DE CONHECIMENTO')
style_run(r, size=22, color=GREEN_DARK, bold=True)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Art · Assistente de Reservas Artezian')
style_run(r2, size=14, color=GREEN_MID)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('artezian.com.br  ·  @ojoaomendonca  ·  Superhost 4.9')
style_run(r3, size=10, color=GRAY_TEXT, italic=True)

doc.add_page_break()

# ============================================================
# 1. ARTEZIAN
# ============================================================
add_h1('1. Sobre a Artezian')
add_body('Empresa de gestão de locação por temporada em Porto Seguro e região, no sul da Bahia. Atendimento direto, sem intermediários como Airbnb ou Booking — mais agilidade e sem taxa de plataforma.')
add_bullet('artezian.com.br', 'Site: ')
add_bullet('@ojoaomendonca — 95 mil seguidores', 'Instagram: ')
add_bullet('Superhost com nota 4.9', 'Classificação: ')
add_bullet('Taperapuã, Mundaí, Arraial D\'Ajuda, Coroa Vermelha', 'Regiões: ')

# ============================================================
# 2. PORTO SEGURO
# ============================================================
add_h1('2. Sobre Porto Seguro')
add_body('Sul da Bahia — praias de água quente o ano inteiro. Voos diretos: SP 45 min · RJ 2h · BH 1h30.')
add_bullet('12 km de faixa com beach clubs (Axé Moi, Tôa Tôa, Boa Beach). Região mais agitada.', 'Taperapuã: ')
add_bullet('Praia tranquila e residencial. Tôa Tôa e Gallo Praia na orla.', 'Mundaí: ')
add_bullet('Vila histórica do outro lado do rio. Acesso por balsa. Mais sossegado.', 'Arraial D\'Ajuda: ')
add_bullet('Praia histórica com recifes de coral naturais.', 'Coroa Vermelha: ')

# ============================================================
# 3. POLÍTICAS
# ============================================================
add_h1('3. Políticas Gerais')

add_h2('Check-in e Checkout')
add_bullet('A partir das 15h', 'Check-in: ')
add_bullet('Até as 12h', 'Checkout: ')
add_bullet('Entrada antecipada ou saída tardia sujeitas à disponibilidade — não garantidas.')

add_h2('Mínimo de Noites')
add_body('3 noites em todos os imóveis, sem exceção.')

add_h2('Taxa de Limpeza Padrão')
add_bullet('R$ 150', 'Imóveis de 1 quarto: ')
add_bullet('R$ 200', 'Imóveis de 2 quartos ou mais: ')
add_bullet('Condomínio do Max: sem taxa de limpeza.')
add_bullet('Casa da Laureana: limpeza diária inclusa na diária.')

add_h2('Animais de Estimação')
add_body('Aceitos apenas no Studio do João (pequeno porte) e em todos os apartamentos do Condomínio do Max, sem taxa adicional. Os demais imóveis não aceitam pets.')

add_h2('Café da Manhã')
add_body('Não incluso. Todos os imóveis têm cozinha equipada (geladeira, fogão, microondas, utensílios).')

add_h2('Eventos e Festas')
add_body('Festas com som alto não são permitidas. Churrascos e confraternizações familiares são bem-vindos, dentro das regras do imóvel.')

# ============================================================
# 4. PAGAMENTO
# ============================================================
add_h1('4. Pagamento')

add_h2('Formas Aceitas')
add_bullet('10% de desconto sobre o valor total', 'Pix à vista: ')
add_bullet('Até 6x com 13% de acréscimo (onde disponível)', 'Cartão de crédito: ')
add_bullet('Pix ou cartão até 3x; acima de 3x requer autorização do gerente', 'Condomínio do Max: ')

add_h2('Como Funciona o Sinal')
add_body('Sinal de 30% a 50% via Pix garante a reserva. O restante é pago no check-in. A porcentagem varia por imóvel (indicada no catálogo).')
add_bullet('50% na assinatura do contrato + 50% no check-in', 'Exceção — Casa da Laureana: ')

add_h2('Contrato')
add_body('Todos os imóveis exigem contrato de locação por temporada, enviado por e-mail antes do pagamento do sinal.')

# ============================================================
# 5. CANCELAMENTO
# ============================================================
add_h1('5. Cancelamento e Reembolso')
add_bullet('Cancelamento com mais de 30 dias: reembolso total do valor pago.')
add_bullet('Cancelamento com 30 dias ou menos: sem reembolso.')
add_bullet('Mudança de datas sujeita à disponibilidade — solicitar com antecedência.')

# ============================================================
# 6. CALENDÁRIO
# ============================================================
add_h1('6. Calendário de Preços')
add_body('Fora dos períodos abaixo, a diária é a de Baixa Temporada. Outubro = Baixa +30%.')

add_h2('2026')
add_table(['Período', 'Tipo'], [
    ['05/01 a 31/01','Alta Temporada'],['05/02 a 09/02','Carnaval'],
    ['09/02 a 14/02','Carnaporto'],['02/04 a 05/04','Feriado'],
    ['18/04 a 22/04','Feriado'],['30/04 a 03/05','Feriado'],
    ['03/06 a 07/06','Feriado'],['07/07 a 24/07','Feriado'],
    ['05/09 a 08/09','Feriado'],['01/10 a 31/10','Alta Temporada (Outubro)'],
    ['10/10 a 18/10','Feriado'],['31/10 a 03/11','Feriado'],
    ['14/11 a 16/11','Feriado'],['19/11 a 22/11','Feriado'],
    ['22/12 a 26/12','Feriado'],['27/12 a 04/01/2027','Réveillon'],
    ['Demais datas','Baixa Temporada'],
])

add_h2('2027')
add_table(['Período', 'Tipo'], [
    ['04/01 a 31/01','Alta Temporada'],['05/02 a 09/02','Carnaval'],
    ['09/02 a 14/02','Carnaporto'],['25/03 a 28/03','Feriado'],
    ['19/04 a 25/04','Feriado'],['30/04 a 02/05','Feriado'],
    ['26/05 a 30/05','Feriado'],['07/07 a 24/07','Feriado'],
    ['06/09 a 08/09','Feriado'],['01/10 a 31/10','Alta Temporada (Outubro)'],
    ['01/11 a 03/11','Feriado'],['13/11 a 16/11','Feriado'],
    ['19/11 a 21/11','Feriado'],['22/12 a 26/12','Feriado'],
    ['27/12 a 04/01/2028','Réveillon'],['Demais datas','Baixa Temporada'],
])

# ============================================================
# 7. CATÁLOGO
# ============================================================
add_h1('7. Catálogo de Imóveis')

# -- Mont Carmelo --
add_h2('Mont Carmelo — Taperapuã')
add_body('400m da praia. Portaria 24h, piscina adulto+infantil, sauna, restaurante, estacionamento.')
add_sep()

add_h3('Studio do João · DS03J')
for k,v in [('Capacidade','Até 5 pessoas'),('Quartos','1 suíte'),('Camas','1 king + 2 solteiros + sofá-cama'),('Banheiros','1'),('Diferenciais','Garagem · aceita pet (pequeno porte)'),('Pagamento','Pix 50% + 50% check-in'),('Taxa de limpeza','R$ 150'),('Link','artezian.com.br/pt/apartment/DS03J')]:
    add_kv(k, v)
add_table(['Período','Diária'],[['Baixa','R$ 250'],['Feriados','R$ 400'],['Alta','R$ 600']])

add_h3('Flat da Mari · DS04J')
for k,v in [('Capacidade','Até 5 pessoas'),('Quartos','1 suíte'),('Camas','1 king + 2 solteiros + sofá-cama'),('Banheiros','2'),('Diferenciais','Térreo com acesso direto à piscina'),('Pagamento','Pix 30% + 70% check-in'),('Taxa de limpeza','R$ 150'),('Link','artezian.com.br/pt/apartment/DS04J')]:
    add_kv(k, v)
add_table(['Período','Diária'],[['Baixa','R$ 250'],['Feriados','R$ 400'],['Alta','R$ 600']])

add_h3('Apto do Emanoel · DS05J')
for k,v in [('Capacidade','Até 8 pessoas'),('Quartos','2 suítes'),('Camas','2 casal + 4 solteiros + sofá-cama'),('Banheiros','3'),('Diferenciais','Churrasqueira privativa · garagem'),('Pagamento','Pix 50% + 50% check-in'),('Taxa de limpeza','R$ 200'),('Link','artezian.com.br/pt/apartment/DS05J')]:
    add_kv(k, v)
add_table(['Período','Diária'],[['Baixa','R$ 250'],['Feriados','R$ 400'],['Alta','R$ 600']])

# -- Varandas --
add_h2('Varandas de Porto — Taperapuã')
add_body('500m da praia. Piscina, churrasqueira coletiva, rampa de acessibilidade. Check-in com a Mara. Roupa de cama inclusa. Não aceita pets.')
add_sep()

add_h3('VP-01 e VP-02 · JR01J / JR02J')
for k,v in [('Capacidade','Até 3 pessoas'),('Quartos','1 suíte'),('Camas','2 camas'),('Banheiros','1'),('Taxa de limpeza','R$ 150')]:
    add_kv(k, v)
add_table(['Período','Diária'],[['Baixa','R$ 190'],['Feriados','R$ 320'],['Janeiro','R$ 500'],['Carnaval','R$ 580'],['Carnaporto','R$ 390'],['Réveillon','R$ 900']])

add_h3('VP-03 e VP-04 · JR03J / JR04J')
for k,v in [('Capacidade','Até 4 pessoas'),('Quartos','1 suíte'),('Camas','2 camas'),('Banheiros','1'),('Taxa de limpeza','R$ 150'),('Preços','Mesmos do VP-01 e VP-02')]:
    add_kv(k, v)

add_h3('VP-05 a VP-08 · JR05J a JR08J')
for k,v in [('Capacidade','Até 5 pessoas'),('Quartos','1 suíte'),('Camas','3 camas'),('Banheiros','1'),('Taxa de limpeza','R$ 150')]:
    add_kv(k, v)
add_table(['Período','Diária'],[['Baixa','R$ 250'],['Feriados','R$ 400'],['Janeiro','R$ 600'],['Carnaval','R$ 580'],['Carnaporto','R$ 450'],['Réveillon','R$ 1.100']])

add_h3('VP-09 · JR09J')
for k,v in [('Capacidade','Até 8 pessoas'),('Quartos','2 suítes'),('Camas','5 camas'),('Banheiros','1 + lavabo'),('Diferenciais','Sacada com vista para a piscina'),('Taxa de limpeza','R$ 200'),('Link','artezian.com.br/pt/apartment/JR09J')]:
    add_kv(k, v)
add_table(['Período','Diária'],[['Baixa','R$ 400'],['Feriados','R$ 600'],['Janeiro','R$ 1.000'],['Carnaval','R$ 880'],['Carnaporto','R$ 650'],['Réveillon','R$ 1.600']])

# -- Max --
add_h2('Condomínio do Max — Mundaí')
add_body('500m da Praia do Mundaí. Piscina, garagem, área de lazer. Ref: Tôa Tôa e Gallo Praia. Por perto: restaurantes, supermercado, farmácia.')
add_body('Check-in 15h (flexível); cofre externo 18h–08h. Roupa de cama inclusa (troca >7 diárias). Sem taxa de limpeza. Aceita pets sem taxa. 14 unidades (8 Tipo 1 + 6 Tipo 2).')
add_sep()

add_h3('Max Tipo 1 — 1 Quarto · 8 unidades')
for k,v in [('Capacidade','Até 5 pessoas'),('Quartos','1'),('Banheiros','1'),('Taxa de limpeza','Não cobrada'),('Pagamento','Pix ou cartão até 3x')]:
    add_kv(k, v)
add_table(['Período','Diária'],[['Baixa','R$ 250'],['Feriados','R$ 400'],['Janeiro','R$ 600'],['Carnaval','R$ 580'],['Carnaporto','R$ 450'],['Réveillon','R$ 1.100']])

add_h3('Max Tipo 2 — 2 Quartos · 6 unidades')
for k,v in [('Capacidade','Até 8 pessoas'),('Quartos','2 (suíte + social)'),('Banheiros','2'),('Diferenciais','Varanda ampla'),('Taxa de limpeza','Não cobrada'),('Pagamento','Pix ou cartão até 3x')]:
    add_kv(k, v)
add_table(['Período','Diária'],[['Baixa','R$ 400'],['Feriados','R$ 600'],['Janeiro','R$ 1.000'],['Carnaval','R$ 880'],['Carnaporto','R$ 650'],['Réveillon','R$ 1.600']])

# -- Avulsos Tap --
add_h2('Apartamentos Avulsos — Taperapuã')
add_sep()

add_h3('Apto da Isa · DS06J')
for k,v in [('Localização','Taperapuã, 400m da praia'),('Capacidade','Até 5 pessoas'),('Quartos','1 suíte'),('Camas','1 king + 2 solteiros + sofá-cama'),('Banheiros','2'),('Diferenciais','Apartamento térreo'),('Taxa de limpeza','R$ 150'),('Link','artezian.com.br/pt/apartment/DS06J')]:
    add_kv(k, v)
add_table(['Período','Diária'],[['Baixa','R$ 250'],['Feriados','R$ 400'],['Alta','R$ 600']])

add_h3('Apto do Reinaldo — Taperapuã · FL10J')
for k,v in [('Localização','Taperapuã, 180m da praia (mais próximo do portfólio)'),('Capacidade','Até 10 pessoas'),('Quartos','3 suítes (1 no térreo)'),('Camas','2 casal + 6 solteiros'),('Banheiros','3'),('Diferenciais','Churrasqueira privativa'),('Pagamento','Pix 30% + 70% check-in'),('Taxa de limpeza','R$ 200'),('Link','artezian.com.br/pt/apartment/FL10J')]:
    add_kv(k, v)
add_table(['Período','Diária'],[['Baixa','R$ 750'],['Feriados','R$ 1.200'],['Alta','R$ 1.600']])

add_h3('Flat da Joyce · HA03J')
for k,v in [('Localização','Taperapuã'),('Capacidade','Até 8 pessoas'),('Quartos','3'),('Camas','2 casal + 4 solteiros + sofá-cama'),('Banheiros','3'),('Diferenciais','Churrasqueira privativa · piscina do condomínio'),('Pagamento','Pix 30% + 70% check-in'),('Taxa de limpeza','R$ 200'),('Link','artezian.com.br/pt/apartment/HA03J')]:
    add_kv(k, v)
add_table(['Período','Diária'],[['Baixa','R$ 650'],['Feriados','R$ 850'],['Alta','R$ 1.200']])

# -- Arraial / Coroa --
add_h2('Arraial D\'Ajuda / Coroa Vermelha')
add_sep()

add_h3('Apto do Reinaldo — Coroa Vermelha · GC01J')
for k,v in [('Localização','Coroa Vermelha'),('Capacidade','Até 12 pessoas'),('Quartos','3 suítes (1 no térreo)'),('Camas','2 casal + 6 solteiros + extras'),('Banheiros','3'),('Diferenciais','Vista mar · churrasqueira privativa · recifes de coral · piscina e sauna do condomínio'),('Pagamento','Pix 50% + 50% check-in'),('Taxa de limpeza','R$ 200'),('Link','artezian.com.br/pt/apartment/GC01J')]:
    add_kv(k, v)
add_table(['Período','Diária'],[['Baixa','R$ 750'],['Feriados','R$ 1.200'],['Alta','R$ 1.600']])

add_h3('Apto da Jessilene · HA02J')
for k,v in [('Localização','Arraial D\'Ajuda'),('Capacidade','Até 12 pessoas'),('Quartos','3 suítes (1 no térreo)'),('Camas','3 casal + 3 solteiros'),('Banheiros','4'),('Diferenciais','Churrasqueira privativa · vila mais tranquila'),('Taxa de limpeza','R$ 200'),('Link','artezian.com.br/pt/apartment/HA02J')]:
    add_kv(k, v)
add_table(['Período','Diária'],[['Baixa','R$ 750'],['Feriados','R$ 1.100'],['Alta','R$ 1.400']])

# -- Casas --
add_h2('Casas')
add_sep()

add_h3('Casa do Tremura · GF02J')
for k,v in [('Localização','Taperapuã, 300m da praia'),('Capacidade','Até 27 pessoas'),('Quartos','6 (4 suítes + 2)'),('Camas','5 casal + 6 solteiros'),('Banheiros','4'),('Diferenciais','Piscina · churrasqueira · melhor custo-benefício para grupos'),('Pagamento','Pix 30% + 70% check-in'),('Link','artezian.com.br/pt/apartment/GF02J')]:
    add_kv(k, v)
add_table(['Período','Diária'],[['Baixa','R$ 1.600'],['Feriados','R$ 2.500'],['Alta','R$ 3.000']])

add_h3('Casa da Laureana · GF04J')
for k,v in [('Localização','Arraial D\'Ajuda'),('Capacidade','Até 13 pessoas'),('Quartos','4 suítes (1 master vista mar)'),('Camas','4 casal + 5 solteiros'),('Banheiros','4'),('Diferenciais','Piscina privativa · limpeza diária inclusa · caução R$ 2.000 · opção mais sofisticada do portfólio'),('Pagamento','50% contrato + 50% check-in'),('Taxa de limpeza','R$ 250 (limpeza diária inclusa)'),('Link','artezian.com.br/pt/apartment/GF04J')]:
    add_kv(k, v)
add_table(['Período','Diária'],[['Baixa','R$ 1.800'],['Feriados','R$ 2.500'],['Alta','R$ 3.000']])

add_h3('Casa da Moana · GF06J')
for k,v in [('Localização','Porto Seguro, 5 min da praia'),('Capacidade','Até 16 pessoas'),('Quartos','5 suítes (3 no térreo)'),('Camas','5 casal + 6 solteiros'),('Banheiros','4'),('Diferenciais','Piscina · forno de lenha · sinuca'),('Pagamento','Pix 50% + 50% check-in'),('Link','artezian.com.br/pt/apartment/GF06J')]:
    add_kv(k, v)
add_table(['Período','Diária'],[['Baixa','R$ 1.500'],['Feriados','R$ 2.500'],['Alta','R$ 3.000']])

add_h3('Casa do John · GG08J')
for k,v in [('Localização','Porto Seguro, 600m da praia'),('Capacidade','Até 31 pessoas'),('Quartos','6'),('Camas','12 beliches + 4 casal + sofá-cama'),('Banheiros','4'),('Diferenciais','Piscina · churrasqueira · freezer · ideal para turmas'),('Pagamento','Pix 30% + 70% check-in'),('Taxa de limpeza','R$ 300'),('Link','artezian.com.br/pt/apartment/GG08J')]:
    add_kv(k, v)
add_table(['Período','Diária'],[['Baixa','R$ 1.800'],['Feriados','R$ 2.200'],['Alta','R$ 2.700']])

add_h3('Casa do Euller · GG06J')
for k,v in [('Localização','Porto Seguro'),('Capacidade','Até 58 pessoas (maior imóvel do portfólio)'),('Quartos','9'),('Camas','9 casal + 19 solteiros'),('Banheiros','9'),('Diferenciais','Piscina grande · churrasqueira externa · área de jogos · sinuca · ideal para eventos e retiros'),('Pagamento','Pix 30% + 70% check-in'),('Link','artezian.com.br/pt/apartment/GG06J')]:
    add_kv(k, v)
add_table(['Período','Diária'],[['Baixa','R$ 2.500'],['Feriados','R$ 3.500'],['Alta','R$ 3.500']])

# ============================================================
# 8. FAQ
# ============================================================
add_h1('8. Perguntas Frequentes')

faqs = [
    ('Pagamento', [
        ('Qual o preço da diária?', 'Varia conforme o imóvel e o período. Informe as datas e o tamanho do grupo.'),
        ('Tem desconto?', 'Pix à vista tem 10% de desconto automático. Preços diretos, sem taxa de plataforma.'),
        ('Posso parcelar no cartão?', 'Sim, até 6x com 13% de acréscimo (onde disponível). No Pix sai mais em conta.'),
        ('Como funciona o pagamento?', 'Sinal de 30%–50% via Pix garante a reserva. Restante pago no check-in.'),
        ('Tem taxa de serviço?', 'Não. Atendimento direto com a Artezian, sem intermediários.'),
    ]),
    ('Reserva', [
        ('Como faço para reservar?', 'Confirme as datas e o grupo pelo WhatsApp. A equipe envia o contrato e o link do sinal.'),
        ('Precisa de contrato?', 'Sim. Enviado por e-mail antes do pagamento.'),
        ('Qual o mínimo de noites?', '3 noites em todos os imóveis.'),
        ('Vocês são do Airbnb?', 'Não. Atendimento direto pela Artezian — mais ágil e sem taxa de plataforma.'),
    ]),
    ('Check-in e Checkout', [
        ('Que horas é o check-in?', 'A partir das 15h. Checkout até as 12h.'),
        ('Posso entrar mais cedo ou sair mais tarde?', 'Sujeito à disponibilidade. Não é garantido previamente.'),
        ('Como funciona o check-in?', 'Mont Carmelo: portaria 24h. Varandas: check-in com a Mara. Max: cofre externo das 18h às 08h.'),
    ]),
    ('Estrutura', [
        ('Tem café da manhã?', 'Não. Todos os imóveis têm cozinha completa equipada.'),
        ('Tem Wi-Fi?', 'Sim, todos os imóveis têm Wi-Fi.'),
        ('Tem ar-condicionado?', 'Sim, todos os imóveis têm ar-condicionado.'),
        ('Tem piscina?', 'Mont Carmelo, Varandas e Max têm piscina de condomínio. Casas têm piscina privativa.'),
        ('Tem limpeza durante a estadia?', 'Laureana tem limpeza diária inclusa. Max faz troca de roupa acima de 7 diárias. Demais: sem limpeza na estadia.'),
    ]),
    ('Pet', [
        ('Quais imóveis aceitam pet?', 'Studio do João (pequeno porte) e Condomínio do Max. Os demais não aceitam.'),
        ('Tem taxa para pet?', 'Não. Sem cobrança adicional nos imóveis que aceitam.'),
    ]),
    ('Cancelamento', [
        ('Posso cancelar?', 'Sim. Acima de 30 dias: reembolso total. Com 30 dias ou menos: sem reembolso.'),
        ('Posso mudar as datas?', 'Sim, mediante disponibilidade. Solicitar com antecedência.'),
    ]),
    ('Localização', [
        ('Qual a distância da praia?', 'Varia por imóvel. Apartamentos: 180m a 600m. Casas: 300m a 5 min de carro.'),
        ('Como chego em Porto Seguro?', 'Aeroporto local. Voos diretos: SP 45 min, RJ 2h, BH 1h30. Aluguel de carro ou transfer.'),
        ('Quais beach clubs ficam perto?', 'Taperapuã: Axé Moi, Tôa Tôa, Boa Beach. Mundaí: Tôa Tôa e Gallo Praia.'),
    ]),
    ('Grupos', [
        ('Tem imóvel para grupo grande?', 'Sim. Portfólio atende de 3 a 58 pessoas. Acima de 30: Casa do John (31p) e Casa do Euller (58p).'),
        ('Pode fazer festa?', 'Festas com som alto não são permitidas. Churrascos e confraternizações familiares, sim.'),
    ]),
]

for sec_title, items in faqs:
    add_h2(sec_title)
    for q, a in items:
        pq = doc.add_paragraph()
        pq.paragraph_format.space_before = Pt(5)
        pq.paragraph_format.space_after  = Pt(1)
        rq = pq.add_run(q)
        style_run(rq, bold=True, color=GREEN_DARK)
        pa = doc.add_paragraph()
        pa.paragraph_format.space_before = Pt(0)
        pa.paragraph_format.space_after  = Pt(6)
        pa.paragraph_format.left_indent  = Cm(0.4)
        ra = pa.add_run(a)
        style_run(ra)

# ============================================================
# SALVAR
# ============================================================
out = r'C:\Users\Mateus Ribas\Desktop\IA\Claude Code\Artezian-RealState\Atendimento IA\base-conhecimento-art.docx'
doc.save(out)
print('Salvo:', out)
